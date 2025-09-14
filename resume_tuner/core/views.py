import io
import os
import google.generativeai as genai
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from PyPDF2 import PdfReader

# Configure Gemini API
genai.configure(api_key="AIzaSyCgM4hoRDSB8ZJOLbGDBVhspg4BTDfcFNg")

@csrf_exempt
def resume_update_view(request):
    if request.method == 'POST':
        # Get resume file (optional)
        resume_file = request.FILES.get('resumeFile')
        # Get job description (mandatory)
        job_description = request.POST.get('jobText')
        job_file = request.POST.get('jobFile')

        if not job_description and not job_file:
            return JsonResponse({"error": "Job description or file is required"}, status=400)

        resume_text = ""

        # Extract resume text if file is provided
        if resume_file:
            filename = resume_file.name
            ext = os.path.splitext(filename)[1].lower()

            try:
                if ext == '.docx':
                    doc = Document(resume_file)
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                elif ext == '.pdf':
                    reader = PdfReader(resume_file)
                    resume_text = "\n".join([page.extract_text() or "" for page in reader.pages])
                else:
                    return JsonResponse({"error": "Only .docx and .pdf files are supported."}, status=400)
            except Exception as e:
                return JsonResponse({"error": f"Error reading resume file: {str(e)}"}, status=500)

        # Gemini prompt
        prompt = f"""
        You are a resume optimization assistant. 
        Here's the original resume (if provided):
        {resume_text}

        Here's the job description:
        {job_description}

        Please rewrite the resume to better align with the job description. Focus on:
        - Highlighting relevant skills and achievements
        - Using action-oriented language
        - Integrating keywords naturally
        - Keeping the format professional

        Return only the updated resume text.
        """

        try:
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(prompt)
            updated_text = response.text
        except Exception as e:
            return JsonResponse({"error": f"Gemini API error: {str(e)}"}, status=500)

        # Create updated DOCX
        updated_doc = Document()
        updated_doc.add_heading('Updated Resume', 0)
        for line in updated_text.split('\n'):
            updated_doc.add_paragraph(line)

        buffer = io.BytesIO()
        updated_doc.save(buffer)
        buffer.seek(0)

        return HttpResponse(
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': 'attachment; filename=updated_resume.docx'}
        )

    return JsonResponse({"error": "Only POST method is allowed"}, status=405)


@csrf_exempt
def prep_schedule_view(request):
    if request.method == 'POST':
        # Optional: resume file
        resume_file = request.FILES.get('resumeFile')
        # Job description text or file
        job_description = request.POST.get('jobText')
        job_file = request.FILES.get('jobFile')

        if not job_description and not job_file:
            return JsonResponse({"error": "Job description or file is required"}, status=400)

        resume_text = ""

        # Extract resume text if file provided
        if resume_file:
            filename = resume_file.name
            ext = os.path.splitext(filename)[1].lower()

            try:
                if ext == '.docx':
                    doc = Document(resume_file)
                    resume_text = "\n".join([para.text for para in doc.paragraphs])
                elif ext == '.pdf':
                    reader = PdfReader(resume_file)
                    resume_text = "\n".join([page.extract_text() or "" for page in reader.pages])
                else:
                    return JsonResponse({"error": "Only .docx and .pdf files are supported."}, status=400)
            except Exception as e:
                return JsonResponse({"error": f"Error reading resume file: {str(e)}"}, status=500)

        # If job description file is provided, extract text
        if job_file:
            filename = job_file.name
            ext = os.path.splitext(filename)[1].lower()
            try:
                if ext == '.docx':
                    doc = Document(job_file)
                    job_description = "\n".join([para.text for para in doc.paragraphs])
                elif ext == '.pdf':
                    reader = PdfReader(job_file)
                    job_description = "\n".join([page.extract_text() or "" for page in reader.pages])
                else:
                    return JsonResponse({"error": "Only .docx and .pdf job files are supported."}, status=400)
            except Exception as e:
                return JsonResponse({"error": f"Error reading job file: {str(e)}"}, status=500)

        # Gemini prompt for schedule
        prompt = f"""
        You are a preparation schedule assistant.
        Here is the resume (if provided):
        {resume_text}

        Here is the job description:
        {job_description}

        Create a **weekly preparation schedule** in HTML table format.
        The schedule should:
        - Include days of the week
        - Include time slots
        - Include specific preparation activities
        - Be visually appealing with inline CSS styling
        - Be ready to render directly in a browser
        """

        try:
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(prompt)
            html_schedule = response.text
        except Exception as e:
            return JsonResponse({"error": f"Gemini API error: {str(e)}"}, status=500)

        return HttpResponse(html_schedule, content_type="text/html")

    return JsonResponse({"error": "Only POST method is allowed"}, status=405)
